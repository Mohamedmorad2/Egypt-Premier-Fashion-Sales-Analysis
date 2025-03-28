%%bash
python - <<'EOF'
from docx import Document
from docx.shared import Pt
import io

# Create a new Document
doc = Document()

# Set the title
doc.add_heading("Town Team Sales Report", level=1)
doc.add_paragraph()

# Analytical questions heading
doc.add_heading("Analytical Questions for Town Team Sales Data:", level=2)
doc.add_paragraph()

# Question 1
doc.add_paragraph("1. What is the best-selling category in terms of quantity?")
doc.add_paragraph("   Best-selling category in terms of value: Women's Wear", style='List Bullet')
doc.add_paragraph("   Total quantity: 678", style='List Bullet')
doc.add_paragraph()

# Question 2
doc.add_paragraph("2. Which region achieved the highest revenue?")
doc.add_paragraph("   Region with the highest total sales: Mansoura", style='List Bullet')
doc.add_paragraph("   Total price: 64,890.15 EGP", style='List Bullet')
doc.add_paragraph()

# Question 3
doc.add_paragraph("3. Which month witnessed the highest sales?")
doc.add_paragraph("   Highest sales month: November", style='List Bullet')
doc.add_paragraph("   Total sales for this month: 24,937.34 EGP", style='List Bullet')
doc.add_paragraph()

# Question 4
doc.add_paragraph("4. Which day witnessed the highest sales?")
doc.add_paragraph("   Highest total sales day: 2024-08-31", style='List Bullet')
doc.add_paragraph("   Total price: 2855.21 EGP", style='List Bullet')
doc.add_paragraph()

# Question 5
doc.add_paragraph("5. Who is the customer who spent the most on purchases?")
doc.add_paragraph("   Highest total sales: C190", style='List Bullet')
doc.add_paragraph("   Total price: 5610.98 EGP", style='List Bullet')
doc.add_paragraph()

# Question 6
doc.add_paragraph("6. Which customer made the most purchases?")
doc.add_paragraph("   The customer with the highest number of purchases is: C099, C190", style='List Bullet')
doc.add_paragraph("   Number of purchases: 13", style='List Bullet')
doc.add_paragraph()

# Question 7
doc.add_paragraph("7. Which product is the most profitable based on total sales?")
doc.add_paragraph("   Most profitable product: Scarf", style='List Bullet')
doc.add_paragraph("   Total price: 5610.98 EGP", style='List Bullet')
doc.add_paragraph()

# Question 8
doc.add_paragraph("8. Which product is the least profitable based on total sales?")
doc.add_paragraph("   Least profitable product: T-Shirt", style='List Bullet')
doc.add_paragraph("   Total price: 73.47 EGP", style='List Bullet')
doc.add_paragraph()

# Question 9
doc.add_paragraph("9. Is there any relationship between the unit price and the quantity sold?")
doc.add_paragraph("   There is no relationship between the unit price and the quantity sold.", style='List Bullet')
doc.add_paragraph()

# Question 10
doc.add_paragraph("10. How were sales distributed among the different categories throughout the year?")
doc.add_paragraph("   a) What is the total sales for each region?", style='List Bullet')
table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Total Sales'
hdr_cells[1].text = 'Region'
row = table.add_row().cells
row[0].text = "64,890.15"
row[1].text = "Mansoura"
row = table.add_row().cells
row[0].text = "63,305.60"
row[1].text = "Alex"
row = table.add_row().cells
row[0].text = "62,583.19"
row[1].text = "Tanta"
row = table.add_row().cells
row[0].text = "49,515.69"
row[1].text = "Madinaty"
row = table.add_row().cells
row[0].text = "8,845.99"
row[1].text = "New Cairo"
row = table.add_row().cells
row[0].text = "5,549.57"
row[1].text = "Nasr City"
doc.add_paragraph()

# Question 11
doc.add_paragraph("11. Are there any seasonal trends in sales?")
doc.add_paragraph("   Yes, there is an increase in sales in May and November.", style='List Bullet')
doc.add_paragraph("   This is because May marks the beginning of summer, leading to higher demand for clothing, and November marks the beginning of winter.", style='List Bullet')
doc.add_paragraph()

# Question 12
doc.add_paragraph("12. What is the percentage contribution of each category to total sales?")
doc.add_paragraph("   Percentage - Category", style='List Bullet')
table2 = doc.add_table(rows=1, cols=2)
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Percentage'
hdr_cells[1].text = 'Category'
row = table2.add_row().cells
row[0].text = "27.7%"
row[1].text = "Women's Wear"
row = table2.add_row().cells
row[0].text = "27.5%"
row[1].text = "Kids' Wear"
row = table2.add_row().cells
row[0].text = "24.1%"
row[1].text = "Men's Wear"
row = table2.add_row().cells
row[0].text = "20.5%"
row[1].text = "Accessories"
doc.add_paragraph()

# Question 13
doc.add_paragraph("13. What is the average order value (Total Sales)?")
doc.add_paragraph("   Minimum sales: 10.83", style='List Bullet')
doc.add_paragraph("   Average sales: 254.69019", style='List Bullet')
doc.add_paragraph("   Maximum sales: 798.34", style='List Bullet')
doc.add_paragraph()

# Question 14
doc.add_paragraph("14. What is the total quantity sold for each product?")
doc.add_paragraph("   Quantity - Product Name", style='List Bullet')
table3 = doc.add_table(rows=1, cols=2)
hdr_cells = table3.rows[0].cells
hdr_cells[0].text = 'Quantity'
hdr_cells[1].text = 'Product Name'
products = [
    ("290", "Polo Shirt"),
    ("290", "Scarf"),
    ("245", "Jeans"),
    ("206", "Socks"),
    ("190", "Belt"),
    ("173", "Gloves"),
    ("151", "Jacket"),
    ("120", "Hat"),
    ("105", "Dress"),
    ("103", "Sweater"),
    ("102", "Shorts"),
    ("98", "Watch"),
    ("93", "Shoes"),
    ("77", "Boots"),
    ("67", "Sandals"),
    ("61", "Skirt"),
    ("56", "Blazer"),
    ("42", "T-Shirt")
]
for qty, name in products:
    row = table3.add_row().cells
    row[0].text = qty
    row[1].text = name
doc.add_paragraph()

# Question 15
doc.add_paragraph("15. Which product generated the highest revenue?")
doc.add_paragraph("   The best-selling product in terms of value is: Scarf", style='List Bullet')
doc.add_paragraph("   Total sales: 30,220.11 EGP", style='List Bullet')
doc.add_paragraph()

# Question 16
doc.add_paragraph("16. How do sales change on a quarterly basis?")
doc.add_paragraph("   There is a significant increase in sales in quarters 3 and 4 compared to quarters 1 and 2.", style='List Bullet')
doc.add_paragraph()

# Question 17
doc.add_paragraph("17. Is there a relationship between the region and the best-selling category within it?")
doc.add_paragraph("   Yes, there is a relationship between the region and the best-selling category, as consumer preferences vary from region to region.", style='List Bullet')
doc.add_paragraph()

# Question 18
doc.add_paragraph("18. What is the percentage of sales coming from repeat customers versus new customers?")
doc.add_paragraph("   New customers: 0.74%", style='List Bullet')
doc.add_paragraph("   Returning customers: 99.2%", style='List Bullet')
doc.add_paragraph()

# Question 19
doc.add_paragraph("19. What are the possible recommendations to improve performance based on the trends and patterns extracted from the data?")
doc.add_paragraph("   Based on the trends and patterns extracted from the data:", style='List Bullet')
doc.add_paragraph("   - Leverage strengths: Focus on the outstanding performance of women's wear, especially in regions with high sales.", style='List Bullet')
doc.add_paragraph("   - Expand the customer base: Develop strategies to attract more new customers to balance the high percentage of repeat buyers.", style='List Bullet')
doc.add_paragraph("   - Seasonal strategies: Optimize inventory and promotional activities during peak sales months (May and November).", style='List Bullet')
doc.add_paragraph("   - Reassess underperforming products: Consider reviewing pricing or marketing strategies for products with weak performance, such as T-Shirts.", style='List Bullet')
doc.add_paragraph()

# Save the document
doc.save("Report Town Team Sales - EN.docx")
print("File 'Report Town Team Sales - EN.docx' created successfully.")
EOF

